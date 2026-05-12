"""Text extraction from PDF and DOCX files.

Returns plaintext + a small metadata dict so downstream agents know
whether the document was OCR'd, how many pages, etc.
"""
import io
import logging
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    """Extract text from a PDF or DOCX file.

    Args:
        file_bytes: raw bytes of the uploaded file.
        filename: original filename, used to detect type via extension.

    Returns:
        (text, metadata)

    Raises:
        ValueError: unsupported file type or parse failure.
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_bytes, filename)
    if suffix == ".docx":
        return _extract_docx(file_bytes, filename)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(file_bytes: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    """Extract text from a PDF. Raises ValueError on parse failure."""
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    # A single unreadable page should not abort the whole document
                    logger.warning("Could not extract text from page %d of %s", page.page_number, filename)
                    pages.append("")
        full_text = "\n\n".join(pages)
    except Exception as exc:
        logger.exception("PDF extraction failed for %s: %s", filename, exc)
        raise ValueError(f"Could not parse PDF '{filename}': {exc}") from exc

    if not full_text.strip():
        logger.warning("PDF '%s' yielded no text — likely a scanned/image-only document", filename)

    return full_text, {
        "type": "pdf",
        "pages": len(pages),
        "size_bytes": len(file_bytes),
        # Suspiciously short text → scanned PDF (no OCR layer)
        "scanned": len(full_text.strip()) < 50 and len(pages) > 0,
    }


def _extract_docx(file_bytes: bytes, filename: str) -> tuple[str, dict[str, Any]]:
    """Extract text from a DOCX. Raises ValueError on parse failure."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        chunks.append(cell.text)
    except Exception as exc:
        logger.exception("DOCX extraction failed for %s: %s", filename, exc)
        raise ValueError(f"Could not parse DOCX '{filename}': {exc}") from exc

    if not chunks:
        logger.warning("DOCX '%s' yielded no text paragraphs", filename)

    return "\n\n".join(chunks), {
        "type": "docx",
        "paragraphs": len(chunks),
        "size_bytes": len(file_bytes),
    }
